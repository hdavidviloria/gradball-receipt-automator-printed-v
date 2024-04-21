import re
from docx import Document

# Open the document template
doc = Document('FINAL_GradBall Layout.docx')

def get_list_of_names(names_file_path):
    items_list = []

    with open(names_file_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line:
                items_list.append(line)
                items_list.append(line)

    return items_list

names_file_path = 'names.txt'

names = get_list_of_names(names_file_path)
name_list = [element.upper() for element in names]

# Define the placeholder pattern
pattern = r'\[NAME\]'

# Loop through the paragraphs in the document
for paragraph in doc.paragraphs:
    # Search for the placeholder pattern in the paragraph text
    if re.search(pattern, paragraph.text):
        # Replace the placeholder with a name from the name list
        name = name_list.pop(0)
        paragraph.text = re.sub(pattern, name, paragraph.text)

# Loop through the tables in the document
for table in doc.tables:
    # Loop through the rows in the table
    for row in table.rows:
        # Loop through the cells in the row
        for cell in row.cells:
            # Loop through the paragraphs in the cell
            for paragraph in cell.paragraphs:
                # Search for the placeholder pattern in the paragraph text
                if re.search(pattern, paragraph.text):
                    # Replace the placeholder with a name from the name list
                    name = name_list.pop(0)
                    paragraph.text = re.sub(pattern, name, paragraph.text)
                    
pattern2 = r'\[LEVEL\]'

def get_list_of_class_section(class_section_file_path):
    items_list = []

    with open(class_section_file_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line:
                items_list.append(line)
                items_list.append(line)

    return items_list

class_section_file_path = 'class_section.txt'


classsection_list = get_list_of_class_section(class_section_file_path)

# Loop through the paragraphs in the document
for paragraph in doc.paragraphs:
    # Search for the placeholder pattern in the paragraph text
    if re.search(pattern2, paragraph.text):
        # Replace the placeholder with a name from the name list
        classsection = classsection_list.pop(0)
        paragraph.text = re.sub(pattern2, classsection, paragraph.text)

# Loop through the tables in the document
for table in doc.tables:
    # Loop through the rows in the table
    for row in table.rows:
        # Loop through the cells in the row
        for cell in row.cells:
            # Loop through the paragraphs in the cell
            for paragraph in cell.paragraphs:
                # Search for the placeholder pattern in the paragraph text
                if re.search(pattern2, paragraph.text):
                    # Replace the placeholder with a name from the name list

                    classsection = classsection_list.pop(0)
                    paragraph.text = re.sub(pattern2, classsection, paragraph.text)

pattern3 = r'\[AR\]'
def generate_number_list(names_file_path):
    arnumber_list = []
    count = 101

    with open(file_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line:
                arnumber_list.append(f'R-{count}')
                arnumber_list.append(f'R-{count}')
                count += 1

    return arnumber_list# Loop through the paragraphs in the document

arnumber_list = generate_number_list(names_file_path)


for paragraph in doc.paragraphs:
    # Search for the placeholder pattern in the paragraph text
    if re.search(pattern3, paragraph.text):
        # Replace the placeholder with a name from the name list
        arnumber = arnumber_list.pop(0)
        paragraph.text = re.sub(pattern3, arnumber, paragraph.text)

# Loop through the tables in the document
for table in doc.tables:
    # Loop through the rows in the table
    for row in table.rows:
        # Loop through the cells in the row
        for cell in row.cells:
            # Loop through the paragraphs in the cell
            for paragraph in cell.paragraphs:
                # Search for the placeholder pattern in the paragraph text
                if re.search(pattern3, paragraph.text):
                    # Replace the placeholder with a name from the name list

                    arnumber = arnumber_list.pop(0)
                    paragraph.text = re.sub(pattern3, arnumber, paragraph.text)


# Save the modified document
doc.save('Modified_Document.docx')